from pathlib import Path
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer
import pdfplumber
from docx import Document
import pandas as pd
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import re
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Dict, List, Set, Tuple, Optional
import json
from datetime import datetime

# ============================================================================
# CONFIGURATION
# ============================================================================

SCORE_THRESHOLD = 0.5
TEXT_CHUNK_SIZE = 2000
POPPLER_PATH = None
EXCLUDED_ENTITIES = {"PERSON", "NRP"}  # Exclude generic Presidio name detection

# Confidence thresholds for PII clustering
CONFIDENCE_HIGH = 0.8
CONFIDENCE_MEDIUM = 0.5
CONFIDENCE_LOW = 0.3

# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class PIIInstance:
    """Represents a single PII detection"""
    entity_type: str
    value: str
    file_path: str
    context: str = ""
    score: float = 0.0
    
    def __hash__(self):
        return hash((self.entity_type, self.value.lower().strip()))
    
    def __eq__(self, other):
        if not isinstance(other, PIIInstance):
            return False
        return (self.entity_type == other.entity_type and 
                self.value.lower().strip() == other.value.lower().strip())

@dataclass
class PIICluster:
    """Represents a group of PII instances that likely belong to the same person"""
    pii_instances: List[PIIInstance] = field(default_factory=list)
    anchor_matches: Set[str] = field(default_factory=set)  # Which anchors matched
    files: Set[str] = field(default_factory=set)
    confidence: float = 0.0
    
    def add_instance(self, instance: PIIInstance):
        self.pii_instances.append(instance)
        self.files.add(instance.file_path)
    
    def get_pii_summary(self) -> Dict[str, int]:
        """Get count of each PII type"""
        summary = defaultdict(int)
        for pii in self.pii_instances:
            summary[pii.entity_type] += 1
        return dict(summary)
    
    def get_unique_pii(self) -> Dict[str, Set[str]]:
        """Get unique values for each PII type"""
        unique = defaultdict(set)
        for pii in self.pii_instances:
            unique[pii.entity_type].add(pii.value)
        return dict(unique)

# ============================================================================
# CUSTOM RECOGNIZERS
# ============================================================================

def create_comprehensive_recognizers():
    """Create all custom PII recognizers (India-specific and general)"""
    recognizers = []
    
    # INDIAN GOVERNMENT IDs
    
    aadhaar_recognizer = PatternRecognizer(
        supported_entity="IN_AADHAAR",
        name="Aadhaar Recognizer",
        patterns=[
            Pattern(name="aadhaar_pattern", regex=r"\b\d{4}\s?\d{4}\s?\d{4}\b", score=0.6)
        ],
        context=["aadhaar", "aadhar", "uid", "uidai"]
    )
    recognizers.append(aadhaar_recognizer)
    
    pan_recognizer = PatternRecognizer(
        supported_entity="IN_PAN",
        name="PAN Recognizer",
        patterns=[
            Pattern(name="pan_pattern", regex=r"\b[A-Z]{5}\d{4}[A-Z]\b", score=0.7)
        ],
        context=["pan", "permanent account number", "income tax"]
    )
    recognizers.append(pan_recognizer)
    
    passport_recognizer = PatternRecognizer(
        supported_entity="IN_PASSPORT",
        name="Indian Passport Recognizer",
        patterns=[
            Pattern(name="passport_pattern", regex=r"\b[A-Z]\d{7}\b", score=0.6)
        ],
        context=["passport", "passport number", "travel document"]
    )
    recognizers.append(passport_recognizer)
    
    voter_id_recognizer = PatternRecognizer(
        supported_entity="IN_VOTER_ID",
        name="Voter ID Recognizer",
        patterns=[
            Pattern(name="voter_id_pattern", regex=r"\b[A-Z]{3}\d{7}\b", score=0.6)
        ],
        context=["voter", "voter id", "epic", "election card"]
    )
    recognizers.append(voter_id_recognizer)
    
    dl_recognizer = PatternRecognizer(
        supported_entity="IN_DRIVING_LICENSE",
        name="Driving License Recognizer",
        patterns=[
            Pattern(name="dl_pattern", regex=r"\b[A-Z]{2}\d{2}\s?\d{11}\b", score=0.6)
        ],
        context=["driving license", "dl", "driver's license", "licence"]
    )
    recognizers.append(dl_recognizer)
    
    vehicle_recognizer = PatternRecognizer(
        supported_entity="IN_VEHICLE_REGISTRATION",
        name="Vehicle Registration Recognizer",
        patterns=[
            Pattern(name="vehicle_pattern", regex=r"\b[A-Z]{2}\s?\d{2}\s?[A-Z]{1,2}\s?\d{4}\b", score=0.6)
        ],
        context=["vehicle", "registration", "rc", "car number", "vehicle number"]
    )
    recognizers.append(vehicle_recognizer)
    
    gstin_recognizer = PatternRecognizer(
        supported_entity="IN_GSTIN",
        name="GSTIN Recognizer",
        patterns=[
            Pattern(name="gstin_pattern", regex=r"\b\d{2}[A-Z]{5}\d{4}[A-Z]\d[Z][A-Z\d]\b", score=0.7)
        ],
        context=["gstin", "gst", "tax", "goods and services tax"]
    )
    recognizers.append(gstin_recognizer)
    
    # CONTACT INFORMATION
    
    indian_phone_recognizer = PatternRecognizer(
        supported_entity="IN_PHONE",
        name="Indian Phone Recognizer",
        patterns=[
            Pattern(name="indian_phone_pattern", regex=r"\b[6-9]\d{9}\b", score=0.5),
            Pattern(name="indian_phone_with_code", regex=r"\+91[\s-]?[6-9]\d{9}\b", score=0.7)
        ],
        context=["phone", "mobile", "contact", "call", "whatsapp", "telephone"]
    )
    recognizers.append(indian_phone_recognizer)
    
    # FINANCIAL DATA
    
    bank_account_recognizer = PatternRecognizer(
        supported_entity="BANK_ACCOUNT",
        name="Bank Account Recognizer",
        patterns=[
            Pattern(name="bank_account_pattern", regex=r"\b\d{9,18}\b", score=0.3)
        ],
        context=["account number", "bank account", "a/c no", "account no", "acc no", "savings account", "current account"]
    )
    recognizers.append(bank_account_recognizer)
    
    ifsc_recognizer = PatternRecognizer(
        supported_entity="IFSC_CODE",
        name="IFSC Code Recognizer",
        patterns=[
            Pattern(name="ifsc_pattern", regex=r"\b[A-Z]{4}0[A-Z0-9]{6}\b", score=0.8)
        ],
        context=["ifsc", "bank code", "branch code", "ifsc code"]
    )
    recognizers.append(ifsc_recognizer)
    
    upi_recognizer = PatternRecognizer(
        supported_entity="UPI_ID",
        name="UPI ID Recognizer",
        patterns=[
            Pattern(name="upi_pattern", regex=r"\b[\w\.-]+@(?:paytm|phonepe|googlepay|gpay|ybl|okaxis|okicici|okhdfcbank|oksbi|axl|ibl|airtel)\b", score=0.7)
        ],
        context=["upi", "upi id", "payment address", "paytm", "phonepe", "gpay"]
    )
    recognizers.append(upi_recognizer)
    
    # PERSONAL INFORMATION
    
    dob_recognizer = PatternRecognizer(
        supported_entity="DATE_OF_BIRTH",
        name="Date of Birth Recognizer",
        patterns=[
            Pattern(name="dob_ddmmyyyy", regex=r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", score=0.3),
            Pattern(name="dob_yyyymmdd", regex=r"\b\d{4}[/-]\d{2}[/-]\d{2}\b", score=0.3)
        ],
        context=["dob", "date of birth", "birth date", "born on", "birthday", "d.o.b"]
    )
    recognizers.append(dob_recognizer)
    
    gender_recognizer = PatternRecognizer(
        supported_entity="GENDER",
        name="Gender Recognizer",
        patterns=[
            Pattern(name="gender_pattern", regex=r"\b(?:male|female|transgender|other|non-binary)\b", score=0.3)
        ],
        context=["gender", "sex", "m/f", "gender identity"]
    )
    recognizers.append(gender_recognizer)
    
    # DIGITAL FOOTPRINTS
    
    ip_recognizer = PatternRecognizer(
        supported_entity="IP_ADDRESS",
        name="IP Address Recognizer",
        patterns=[
            Pattern(name="ipv4_pattern", regex=r"\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b", score=0.6)
        ],
        context=["ip address", "ipv4", "server ip", "host"]
    )
    recognizers.append(ip_recognizer)
    
    mac_recognizer = PatternRecognizer(
        supported_entity="MAC_ADDRESS",
        name="MAC Address Recognizer",
        patterns=[
            Pattern(name="mac_pattern", regex=r"\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b", score=0.8)
        ],
        context=["mac address", "hardware address", "device mac"]
    )
    recognizers.append(mac_recognizer)
    
    # SENSITIVE DATA
    
    medical_record_recognizer = PatternRecognizer(
        supported_entity="MEDICAL_RECORD",
        name="Medical Record Recognizer",
        patterns=[
            Pattern(name="mrn_pattern", regex=r"\b(?:MRN|MR|PRN)[\s:-]?[A-Z0-9]{6,10}\b", score=0.6)
        ],
        context=["medical record", "patient id", "hospital record", "mrn", "patient number"]
    )
    recognizers.append(medical_record_recognizer)
    
    blood_group_recognizer = PatternRecognizer(
        supported_entity="BLOOD_GROUP",
        name="Blood Group Recognizer",
        patterns=[
            Pattern(name="blood_group_pattern", regex=r"\b(?:A|B|AB|O)[+-]\s*(?:ve|ive|positive|negative)?\b", score=0.6)
        ],
        context=["blood group", "blood type", "blood"]
    )
    recognizers.append(blood_group_recognizer)
    
    return recognizers

# ============================================================================
# ANCHOR DETECTION
# ============================================================================

def create_anchor_recognizers(anchors: Dict[str, str]) -> List[PatternRecognizer]:
    """Create custom recognizers for user-provided anchor values"""
    recognizers = []
    
    for anchor_type, anchor_value in anchors.items():
        if not anchor_value:
            continue
            
        if anchor_type == "name":
            # Exact name matching with word boundaries
            escaped_name = re.escape(anchor_value)
            name_pattern = r'\b' + escaped_name + r'\b'
            
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_NAME",
                name="Anchor Name Recognizer",
                patterns=[
                    Pattern(name="anchor_name_pattern", regex=name_pattern, score=0.95)
                ],
                context=[]
            )
            recognizers.append(recognizer)
            
        elif anchor_type == "email":
            # Email anchor
            escaped_email = re.escape(anchor_value)
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_EMAIL",
                name="Anchor Email Recognizer",
                patterns=[
                    Pattern(name="anchor_email_pattern", regex=escaped_email, score=0.95)
                ],
                context=["email", "mail"]
            )
            recognizers.append(recognizer)
            
        elif anchor_type == "phone":
            # Phone anchor (handle various formats)
            clean_phone = re.sub(r'[\s\-\(\)]', '', anchor_value)
            phone_pattern = r'\b' + re.escape(clean_phone) + r'\b'
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_PHONE",
                name="Anchor Phone Recognizer",
                patterns=[
                    Pattern(name="anchor_phone_pattern", regex=phone_pattern, score=0.95)
                ],
                context=["phone", "mobile", "contact"]
            )
            recognizers.append(recognizer)
            
        elif anchor_type == "aadhaar":
            # Aadhaar anchor
            clean_aadhaar = re.sub(r'\s', '', anchor_value)
            # Allow flexible spacing in matching
            if len(clean_aadhaar) == 12:
                pattern = r'\b' + clean_aadhaar[0:4] + r'\s?' + clean_aadhaar[4:8] + r'\s?' + clean_aadhaar[8:12] + r'\b'
                recognizer = PatternRecognizer(
                    supported_entity="ANCHOR_AADHAAR",
                    name="Anchor Aadhaar Recognizer",
                    patterns=[
                        Pattern(name="anchor_aadhaar_pattern", regex=pattern, score=0.95)
                    ],
                    context=["aadhaar", "aadhar", "uid"]
                )
                recognizers.append(recognizer)
                
        elif anchor_type == "pan":
            # PAN anchor
            pan_pattern = r'\b' + re.escape(anchor_value.upper()) + r'\b'
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_PAN",
                name="Anchor PAN Recognizer",
                patterns=[
                    Pattern(name="anchor_pan_pattern", regex=pan_pattern, score=0.95)
                ],
                context=["pan", "permanent account"]
            )
            recognizers.append(recognizer)
    
    return recognizers

def check_anchors_in_text(text: str, anchors: Dict[str, str]) -> Set[str]:
    """Check which anchors appear in the text (case-insensitive, flexible matching)"""
    matched_anchors = set()
    text_lower = text.lower()
    
    for anchor_type, anchor_value in anchors.items():
        if not anchor_value:
            continue
            
        if anchor_type == "name":
            # Normalize and check name
            name_normalized = ' '.join(anchor_value.lower().strip().split())
            if name_normalized in text_lower:
                matched_anchors.add(f"name:{anchor_value}")
                
        elif anchor_type == "email":
            if anchor_value.lower() in text_lower:
                matched_anchors.add(f"email:{anchor_value}")
                
        elif anchor_type == "phone":
            # Remove formatting for phone comparison
            clean_phone = re.sub(r'[\s\-\(\)]', '', anchor_value)
            clean_text = re.sub(r'[\s\-\(\)]', '', text_lower)
            if clean_phone.lower() in clean_text:
                matched_anchors.add(f"phone:{anchor_value}")
                
        elif anchor_type == "aadhaar":
            # Check with and without spaces
            clean_aadhaar = re.sub(r'\s', '', anchor_value)
            clean_text = re.sub(r'\s', '', text_lower)
            if clean_aadhaar in clean_text:
                matched_anchors.add(f"aadhaar:{anchor_value}")
                
        elif anchor_type == "pan":
            if anchor_value.upper() in text.upper():
                matched_anchors.add(f"pan:{anchor_value}")
    
    return matched_anchors

# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def chunk_text(text, size=TEXT_CHUNK_SIZE):
    """Split text into chunks for processing"""
    for i in range(0, len(text), size):
        yield text[i:i + size]

def infer_column_context(column_name: str) -> List[str]:
    """Infer semantic context from column name"""
    col = column_name.lower()
    context_keywords = []
    
    if any(k in col for k in ["name", "student", "employee", "candidate", "person", "patient"]):
        context_keywords.extend(["name", "person"])
    
    if any(k in col for k in ["phone", "mobile", "contact", "tel"]):
        context_keywords.extend(["phone", "contact"])
    
    if "email" in col or "mail" in col:
        context_keywords.append("email")
    
    if any(k in col for k in ["address", "location", "residence", "street"]):
        context_keywords.extend(["address", "location"])
    
    if any(k in col for k in ["dob", "birth"]):
        context_keywords.extend(["date of birth", "dob"])
    
    if "gender" in col or "sex" in col:
        context_keywords.append("gender")
    
    if any(k in col for k in ["aadhaar", "aadhar"]):
        context_keywords.extend(["aadhaar", "uid"])
    
    if "pan" in col and "company" not in col:
        context_keywords.extend(["pan", "permanent account number"])
    
    if any(k in col for k in ["passport"]):
        context_keywords.append("passport")
    
    if any(k in col for k in ["vehicle", "registration", "car", "rc"]):
        context_keywords.extend(["vehicle", "registration"])
    
    if any(k in col for k in ["voter", "epic"]):
        context_keywords.extend(["voter", "voter id"])
    
    if any(k in col for k in ["license", "licence", "dl"]):
        context_keywords.extend(["driving license", "dl"])
    
    if any(k in col for k in ["gstin", "gst"]):
        context_keywords.extend(["gstin", "gst"])
    
    if any(k in col for k in ["account"]) and any(k in col for k in ["bank", "savings", "current"]):
        context_keywords.extend(["bank account", "account number"])
    
    if any(k in col for k in ["ifsc"]):
        context_keywords.extend(["ifsc", "bank code"])
    
    if any(k in col for k in ["upi"]):
        context_keywords.extend(["upi", "payment address"])
    
    if any(k in col for k in ["medical", "patient", "mrn"]):
        context_keywords.extend(["medical record", "patient id"])
    
    if any(k in col for k in ["blood"]):
        context_keywords.append("blood group")
    
    return context_keywords

def create_semantic_context(column_name: str, value: str) -> str:
    """Create semantic context for better PII detection"""
    context_keywords = infer_column_context(column_name)
    
    if context_keywords:
        context_str = ", ".join(context_keywords)
        return f"{context_str}: {value}. The column '{column_name}' contains the value {value}."
    else:
        return f"The column '{column_name}' contains the value {value}."

def extract_text_from_file(file_path: Path, anchors: Dict[str, str]) -> Tuple[bool, List[Tuple[str, Set[str]]]]:
    """
    Extract text from file and check for anchor matches.
    Returns: (has_anchors, [(text_chunk, matched_anchors), ...])
    """
    suffix = file_path.suffix.lower()
    text_chunks = []
    has_any_anchor = False

    try:
        if suffix == ".txt":
            content = file_path.read_text(errors="ignore")
            matched = check_anchors_in_text(content, anchors)
            if matched:
                has_any_anchor = True
            for chunk in chunk_text(content):
                chunk_matched = check_anchors_in_text(chunk, anchors)
                text_chunks.append((chunk, chunk_matched))

        elif suffix == ".pdf":
            all_text = []
            pages_with_text = []
            pages_without_text = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        all_text.append(text)
                        pages_with_text.append(page_num)
                        matched = check_anchors_in_text(text, anchors)
                        if matched:
                            has_any_anchor = True
                    else:
                        pages_without_text.append(page_num)
            
            # Process text pages
            for text in all_text:
                for chunk in chunk_text(text):
                    chunk_matched = check_anchors_in_text(chunk, anchors)
                    text_chunks.append((chunk, chunk_matched))
            
            # OCR for image-based pages
            if pages_without_text:
                try:
                    kwargs = {'dpi': 300, 'fmt': 'png'}
                    if POPPLER_PATH:
                        kwargs['poppler_path'] = POPPLER_PATH
                        
                    images = convert_from_path(file_path, **kwargs)
                    custom_config = r'--oem 3 --psm 6'
                    
                    for page_num, img in enumerate(images, start=1):
                        if page_num in pages_without_text:
                            try:
                                page_text = pytesseract.image_to_string(img, lang='eng', config=custom_config)
                                if page_text.strip():
                                    matched = check_anchors_in_text(page_text, anchors)
                                    if matched:
                                        has_any_anchor = True
                                    for chunk in chunk_text(page_text):
                                        chunk_matched = check_anchors_in_text(chunk, anchors)
                                        text_chunks.append((chunk, chunk_matched))
                            except Exception:
                                continue
                except Exception:
                    pass

        elif suffix == ".docx":
            doc = Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            matched = check_anchors_in_text(full_text, anchors)
            if matched:
                has_any_anchor = True
            
            for para in doc.paragraphs:
                if para.text.strip():
                    para_matched = check_anchors_in_text(para.text, anchors)
                    text_chunks.append((para.text, para_matched))

        elif suffix.startswith(".xls"):
            try:
                sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
            except Exception:
                return False, []

            for sheet_name, df in sheets.items():
                df = df.fillna("")
                
                sheet_text = df.to_string()
                matched = check_anchors_in_text(sheet_text, anchors)
                if matched:
                    has_any_anchor = True

                for _, row in df.iterrows():
                    row_sentences = []
                    for col in df.columns:
                        value = str(row[col]).strip()
                        if value:
                            row_sentences.append(create_semantic_context(col, value))

                    if row_sentences:
                        row_text = " ".join(row_sentences)
                        row_matched = check_anchors_in_text(row_text, anchors)
                        for chunk in chunk_text(row_text):
                            text_chunks.append((chunk, row_matched))

        elif suffix == ".csv":
            full_df = pd.read_csv(file_path, dtype=str)
            full_text = full_df.to_string()
            
            matched = check_anchors_in_text(full_text, anchors)
            if matched:
                has_any_anchor = True
            
            for chunk_df in pd.read_csv(file_path, chunksize=500, dtype=str):
                chunk_df = chunk_df.fillna("")
                for _, row in chunk_df.iterrows():
                    row_sentences = []
                    for col in chunk_df.columns:
                        value = str(row[col]).strip()
                        if value:
                            row_sentences.append(create_semantic_context(col, value))
                    
                    if row_sentences:
                        row_text = " ".join(row_sentences)
                        row_matched = check_anchors_in_text(row_text, anchors)
                        text_chunks.append((row_text, row_matched))

    except Exception as e:
        return False, []

    return has_any_anchor, text_chunks

# ============================================================================
# PII ANALYSIS
# ============================================================================

def analyze_text_for_pii(text: str, analyzer: AnalyzerEngine, file_path: str, 
                         matched_anchors: Set[str]) -> List[PIIInstance]:
    """Analyze text chunk and return PII instances"""
    instances = []
    
    try:
        results = analyzer.analyze(
            text=text,
            language="en",
            score_threshold=SCORE_THRESHOLD
        )
        
        for result in results:
            entity_type = result.entity_type
            
            # Skip excluded entities
            if entity_type in EXCLUDED_ENTITIES:
                continue
            
            # Extract the actual value
            value = text[result.start:result.end]
            
            # Create context snippet (50 chars before and after)
            context_start = max(0, result.start - 50)
            context_end = min(len(text), result.end + 50)
            context = text[context_start:context_end].replace('\n', ' ').strip()
            
            instance = PIIInstance(
                entity_type=entity_type,
                value=value,
                file_path=file_path,
                context=context,
                score=result.score
            )
            instances.append(instance)
            
    except Exception as e:
        pass
    
    return instances

# ============================================================================
# CLUSTERING & CONFIDENCE SCORING
# ============================================================================

def calculate_cluster_confidence(cluster: PIICluster, total_files_scanned: int) -> float:
    """
    Calculate confidence score for a PII cluster based on:
    1. Number of anchor matches (higher = better)
    2. Number of strong PII types (Aadhaar, PAN, etc.)
    3. Co-occurrence patterns
    4. File distribution
    """
    confidence = 0.0
    
    # Factor 1: Anchor matches (0-40 points)
    num_anchors = len(cluster.anchor_matches)
    if num_anchors >= 3:
        confidence += 40
    elif num_anchors == 2:
        confidence += 30
    elif num_anchors == 1:
        confidence += 20
    
    # Factor 2: Strong identifiers (0-30 points)
    strong_identifiers = {
        'IN_AADHAAR', 'IN_PAN', 'IN_PASSPORT', 'IN_VOTER_ID', 
        'IN_DRIVING_LICENSE', 'ANCHOR_AADHAAR', 'ANCHOR_PAN'
    }
    unique_pii = cluster.get_unique_pii()
    strong_count = sum(1 for entity_type in unique_pii.keys() if entity_type in strong_identifiers)
    confidence += min(strong_count * 10, 30)
    
    # Factor 3: Diversity of PII types (0-20 points)
    num_pii_types = len(unique_pii)
    confidence += min(num_pii_types * 3, 20)
    
    # Factor 4: Multiple files (0-10 points)
    num_files = len(cluster.files)
    if num_files > 1:
        confidence += min(num_files * 3, 10)
    
    # Normalize to 0-1 scale
    return min(confidence / 100.0, 1.0)

def create_pii_cluster(pii_instances: List[PIIInstance], matched_anchors: Set[str], 
                       total_files: int) -> PIICluster:
    """Create a PII cluster from instances"""
    cluster = PIICluster()
    cluster.anchor_matches = matched_anchors
    
    for instance in pii_instances:
        cluster.add_instance(instance)
    
    cluster.confidence = calculate_cluster_confidence(cluster, total_files)
    
    return cluster

# ============================================================================
# MAIN DISCOVERY ENGINE
# ============================================================================

def discover_pii(folder_path: str, **anchors) -> Dict[str, any]:
    """
    Main PII discovery function with clustering.
    
    Args:
        folder_path: Path to scan
        **anchors: Flexible keyword arguments for anchors
            - name: Person's full name
            - email: Email address
            - phone: Phone number
            - aadhaar: Aadhaar number
            - pan: PAN number
    
    Returns:
        Dictionary with discovery results and clustering analysis
    """
    
    # Validate inputs
    if not anchors:
        raise ValueError("At least one anchor must be provided (name, email, phone, aadhaar, or pan)")
    
    root = Path(folder_path)
    if not root.exists():
        raise ValueError(f"Folder path does not exist: {folder_path}")
    
    # Initialize analyzer with all recognizers
    analyzer = AnalyzerEngine()
    
    # Add comprehensive recognizers
    comprehensive_recognizers = create_comprehensive_recognizers()
    for recognizer in comprehensive_recognizers:
        analyzer.registry.add_recognizer(recognizer)
    
    # Add anchor-specific recognizers
    anchor_recognizers = create_anchor_recognizers(anchors)
    for recognizer in anchor_recognizers:
        analyzer.registry.add_recognizer(recognizer)
    
    print("\n" + "="*80)
    print("PII DISCOVERY")
    print("="*80)
    print(f"\nScanning: {folder_path}\n")
    
    # Data structures for results
    all_pii_instances = []
    files_scanned = 0
    files_with_anchors = 0
    files_with_pii = 0
    global_anchor_matches = set()
    
    # PASS 1: Scan all files
    print("Scanning files...")
    
    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue
        
        files_scanned += 1
        
        # Extract text and check for anchors
        has_anchors, text_chunks = extract_text_from_file(file_path, anchors)
        
        if not has_anchors:
            continue
        
        files_with_anchors += 1
        file_has_pii = False
        
        # Analyze each chunk for PII
        for text_chunk, chunk_anchors in text_chunks:
            if not text_chunk.strip():
                continue
            
            # Track which anchors were found
            global_anchor_matches.update(chunk_anchors)
            
            # Extract PII from this chunk
            pii_instances = analyze_text_for_pii(
                text_chunk, 
                analyzer, 
                str(file_path),
                chunk_anchors
            )
            
            if pii_instances:
                file_has_pii = True
                all_pii_instances.extend(pii_instances)
        
        if file_has_pii:
            files_with_pii += 1
    
    print(f"Scanned {files_scanned} files.\n")
    
    cluster = create_pii_cluster(all_pii_instances, global_anchor_matches, files_scanned)
    
    # Prepare results
    results = {
        'search_timestamp': datetime.now().isoformat(),
        'search_parameters': {k: v for k, v in anchors.items() if v},
        'statistics': {
            'total_files_scanned': files_scanned,
            'files_with_anchors': files_with_anchors,
            'files_with_pii': files_with_pii,
            'total_pii_instances': len(all_pii_instances),
            'unique_files_with_pii': len(cluster.files)
        },
        'cluster': {
            'confidence': cluster.confidence,
            'confidence_level': (
                'HIGH' if cluster.confidence >= CONFIDENCE_HIGH else
                'MEDIUM' if cluster.confidence >= CONFIDENCE_MEDIUM else
                'LOW'
            ),
            'matched_anchors': list(cluster.anchor_matches),
            'pii_summary': cluster.get_pii_summary(),
            'unique_pii_values': {k: list(v) for k, v in cluster.get_unique_pii().items()},
            'files': list(cluster.files)
        },
        'detailed_instances': []
    }
    
    # Add detailed instances (grouped by file)
    instances_by_file = defaultdict(list)
    for instance in all_pii_instances:
        instances_by_file[instance.file_path].append(instance)
    
    for file_path, instances in instances_by_file.items():
        file_data = {
            'file': file_path,
            'pii_count': len(instances),
            'pii_types': defaultdict(list)
        }
        
        for instance in instances:
            file_data['pii_types'][instance.entity_type].append({
                'value': instance.value,
                'context': instance.context,
                'score': instance.score
            })
        
        file_data['pii_types'] = dict(file_data['pii_types'])
        results['detailed_instances'].append(file_data)
    
    return results

# ============================================================================
# REPORTING
# ============================================================================

def print_discovery_report(results: Dict):
    """Print a simplified discovery report - only file paths and PII types"""
    
    print("\n" + "="*80)
    print("PII DISCOVERY REPORT")
    print("="*80 + "\n")
    
    # Get detailed instances grouped by file
    detailed_instances = results.get('detailed_instances', [])
    
    if not detailed_instances:
        print("No PII found for the specified person.")
        print("\n" + "="*80 + "\n")
        return
    
    print(f"Found PII in {len(detailed_instances)} file(s):\n")
    
    # Print each file with its PII types
    for file_data in detailed_instances:
        file_path = file_data['file']
        pii_types = file_data['pii_types']
        
        print(f"FILE: {file_path}")
        
        # List unique PII types found in this file
        for pii_type in sorted(pii_types.keys()):
            print(f"  - {pii_type}")
        
        print()  # Empty line between files
    
    print("="*80 + "\n")

def save_report_json(results: Dict, output_path: str):
    """Save detailed report as JSON"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"Detailed report saved to: {output_path}")

# ============================================================================
# CLI INTERFACE
# ============================================================================

def main():
    """Interactive CLI for PII discovery"""
    
    print("\n" + "="*80)
    print("PII DISCOVERY SYSTEM")
    print("="*80 + "\n")
    
    print("Provide at least ONE identifier to search for the person's PII.\n")
    
    # Collect anchors
    anchors = {}
    
    name = input("Person's full name (optional, press Enter to skip): ").strip()
    if name:
        anchors['name'] = name
    
    email = input("Email address (optional, press Enter to skip): ").strip()
    if email:
        anchors['email'] = email
    
    phone = input("Phone number (optional, press Enter to skip): ").strip()
    if phone:
        anchors['phone'] = phone
    
    aadhaar = input("Aadhaar number (optional, press Enter to skip): ").strip()
    if aadhaar:
        anchors['aadhaar'] = aadhaar
    
    pan = input("PAN number (optional, press Enter to skip): ").strip()
    if pan:
        anchors['pan'] = pan
    
    if not anchors:
        print("\n❌ ERROR: At least one identifier must be provided.")
        return
    
    folder = input("\nFolder path to scan: ").strip()
    
    if not folder or not Path(folder).exists():
        print(f"\n❌ ERROR: Invalid folder path: {folder}")
        return
    
    # Run discovery
    try:
        results = discover_pii(folder, **anchors)
        
        # Print report
        print_discovery_report(results)
        
        # Ask to save JSON
        save_json = input("Save detailed report as JSON? (y/n): ").strip().lower()
        if save_json == 'y':
            output_path = input("Output file path (e.g., report.json): ").strip()
            if output_path:
                save_report_json(results, output_path)
        
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
